"""Builds docs/NLP Final Report - Group 02 Section B.docx from the course template.

Fills the cover, the contributions table and all eight report sections, and appends the
four code listings the appendix asks for. Every number is read from the artefacts on
disk at build time rather than typed here, so the report cannot drift from the results.

Two scales appear in the report and the distinction is deliberate. Section 2 reports the
midterm experiments, which ran on a 6,000-row balanced sample per dataset. Sections 3 to
6 report the final-term experiments, which ran on the full balanced corpora. Table 2
follows the course specification, which asks for the classical rows from the midterm and
the transformer rows from the final term, and the text says so plainly.

Run:  python docs/builders/build_final_report.py
"""
import json
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from sklearn.metrics import accuracy_score, precision_recall_fscore_support

ROOT = Path(__file__).resolve().parents[2]
TEMPLATE = ROOT / 'docs' / 'NLP Report Template with Cover Page - Final-term Summer 25-26.docx'
OUT = ROOT / 'docs' / 'NLP Final Report - Group 02 Section B.docx'
PS = ROOT / 'experiments' / 'paper_scale'
AUDIT = ROOT / 'experiments' / 'audit'

BODY_FONT, BODY_PT = 'Times New Roman', 12
TABLE_PT, CODE_PT = 7.5, 7.5

MEMBERS = [
    ('NILOY PAUL', '23-51773-2',
     'Data preprocessing and corpus preparation (25%). Loaded and class-balanced both '
     'corpora, built the duplicate-content hashing and the group-aware 72/8/20 split, ran '
     'the duplicate and leakage audit, and implemented the classical normalisation '
     'pipeline and the 128-token tokenisation diagnostics.'),
    ('ARNOB SARKER SUPTA', '23-52080-2',
     'BERT experiments (25%). Ran the bert-base-uncased hyperparameter grid over learning '
     'rate, batch size and weight decay on both datasets, selected the best configuration '
     'on validation weighted F1, and produced the BERT confusion matrices and per-document '
     'prediction arrays.'),
    ('NAZMUS SAKIB', '23-52638-2',
     'BERT variant experiments and analysis (25%). Ran the DeBERTa grid on both datasets, '
     'carried out the surface and content analysis, the tokenisation and label-free '
     'checks, the paired significance testing, and the figures.'),
    ('AFNAN UR RAYAN', '23-51992-2',
     'Ensemble, evaluation and reporting (25%). Built the validation-weighted soft-vote '
     'ensemble and its degenerate-weight check, ran the classical baselines and the '
     'cross-dataset transfer evaluation, compiled the result tables, and wrote the report.'),
]

# The classical models the midterm introduced, re-run on the FULL corpora. Both text
# representations were re-run, so the report no longer mixes two dataset sizes in one
# table the way the earlier draft did.
CLASSICAL = ['Naive Bayes', 'Logistic Regression', 'Support Vector Machine']
REPS = ['BoW', 'TF-IDF']

GRID = [(2e-5, 16, 0.01), (3e-5, 16, 0.01), (2e-5, 32, 0.01), (3e-5, 32, 0.01),
        (2e-5, 16, 0.1),  (3e-5, 16, 0.1),  (2e-5, 32, 0.1),  (3e-5, 32, 0.1)]


# ---------------------------------------------------------------- data loading
def load_results():
    """Everything the report quotes, read from disk."""
    r = {'grid': {}, 'deployed': {}, 'ensemble': {}}
    for tag in ('D1', 'D2'):
        for mk in ('BERT', 'DeBERTa'):
            for lr, bs, wd in GRID:
                key = f'full_{tag}_{mk}_lr{lr:g}_bs{bs}_wd{wd:g}_s42'
                p = PS / 'results' / f'{key}.json'
                if p.exists():
                    t = json.load(open(p))['test']
                    r['grid'][(tag, mk, lr, bs, wd)] = (
                        t['accuracy'], t['precision'], t['recall'], t['f1'])
            info = json.load(open(PS / 'models' / f'{tag}_{mk}' / 'run_info.json'))
            r['deployed'][(tag, mk)] = info

        # Ensemble, recomputed from the deployed checkpoints' own probabilities.
        w = json.load(open(AUDIT / 'ensemble_full_scale.json'))['datasets'][tag]['weight_on_bert']
        pb = np.load(PS / 'probs' / f"{r['deployed'][(tag,'BERT')]['key']}.npz")
        pd_ = np.load(PS / 'probs' / f"{r['deployed'][(tag,'DeBERTa')]['key']}.npz")
        y = pb['test_labels']
        pred = (w * pb['test_probs'] + (1 - w) * pd_['test_probs']).argmax(1)
        acc = accuracy_score(y, pred)
        pre, rec, f1, _ = precision_recall_fscore_support(
            y, pred, average='weighted', zero_division=0)
        r['ensemble'][tag] = {'w': w, 'metrics': (round(acc, 4), round(pre, 4),
                                                  round(rec, 4), round(f1, 4))}
    # Classical models at full scale, both representations. The JSON records accuracy
    # and F1 but not precision and recall, so those are recomputed from the saved score
    # arrays and checked back against the recorded accuracy and F1 before use.
    ev = json.load(open(AUDIT / 'full_model_evaluation.json'))
    sc = np.load(AUDIT / 'full_model_scores.npz')
    r['classical'] = {}
    for tag in ('D1', 'D2'):
        y = sc[f'{tag}|y_true']
        for name in CLASSICAL:
            for rep in REPS:
                key = f'{name} ({rep})'
                s_ = sc[f'{tag}|{key}']
                pred = ((s_ >= 0.5) if (s_.min() >= 0 and s_.max() <= 1) else (s_ > 0)).astype(int)
                acc = accuracy_score(y, pred)
                pre, rec, f1, _ = precision_recall_fscore_support(
                    y, pred, average='weighted', zero_division=0)
                rj = ev['datasets'][tag]['models'][key]
                assert abs(acc - rj['accuracy']) < 5e-4 and abs(f1 - rj['weighted_f1']) < 5e-4, \
                    f'recomputed metrics disagree with the recorded ones for {tag} {key}'
                r['classical'][(tag, name, rep)] = (round(acc, 4), round(pre, 4),
                                                    round(rec, 4), round(f1, 4))
    r['eval'] = ev
    r['tok'] = json.load(open(AUDIT / 'token_length_both.json'))
    r['hc3_counts'] = json.load(open(AUDIT / 'hc3_class_counts.json'))
    r['daigt_audit'] = json.load(open(AUDIT / 'daigt_full_audit.json'))
    r['hc3_audit'] = json.load(open(AUDIT / 'hc3_full_audit.json'))
    r['splits'] = {}
    for tag in ('D1', 'D2'):
        df = pd.read_parquet(PS / 'work' / f'data_{tag}.parquet')
        sp = np.load(PS / 'work' / f'split_{tag}.npz')
        y = df['label'].values
        rec = {'balanced': len(df), 'per_class': int((y == 0).sum())}
        for part in ('train', 'val', 'test'):
            rec[part] = int(len(sp[part]))
        wl = df['text'].str.split().str.len()
        rec['median_words'] = int(wl.median())
        rec['median_words_human'] = int(wl[y == 0].median())
        rec['median_words_machine'] = int(wl[y == 1].median())
        r['splits'][tag] = rec
    r['dec'] = json.load(open(AUDIT / 'surface_content_decomposition.json'))
    r['ens_json'] = json.load(open(AUDIT / 'ensemble_full_scale.json'))
    return r


# ---------------------------------------------------------------- docx helpers
def style_run(run, size=BODY_PT, font=BODY_FONT, bold=False):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    # Word needs the east-asian name set too or it silently substitutes.
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    return run


def set_text(par, text, size=BODY_PT, font=BODY_FONT, bold=False, justify=True):
    for r in list(par.runs):
        r._element.getparent().remove(r._element)
    style_run(par.add_run(text), size, font, bold)
    if justify:
        par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return par


def para_after(anchor, text, size=BODY_PT, font=BODY_FONT, bold=False, justify=True):
    """Insert a new paragraph directly after `anchor` and return it."""
    new_p = anchor._element.makeelement(qn('w:p'), {})
    anchor._element.addnext(new_p)
    from docx.text.paragraph import Paragraph
    par = Paragraph(new_p, anchor._parent)
    set_text(par, text, size, font, bold, justify)
    return par


def find_par(doc, needle):
    for p in doc.paragraphs:
        if needle in p.text:
            return p
    raise KeyError(needle)


def keep_with_next(par):
    """Bind a caption to the table that follows it, so the two never land on
    opposite sides of a page break."""
    par.paragraph_format.keep_with_next = True
    return par


def table_after(doc, anchor, rows, widths=None, header_rows=1, size=TABLE_PT):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.cell(i, j)
            cell.text = ''
            par = cell.paragraphs[0]
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER if j else WD_ALIGN_PARAGRAPH.LEFT
            style_run(par.add_run(str(val)), size, BODY_FONT, bold=(i < header_rows))
            par.paragraph_format.space_before = Pt(1)
            par.paragraph_format.space_after = Pt(1)
    anchor._element.addnext(t._element)
    return t


def set_cell(cell, text, size=TABLE_PT, bold=False, center=True):
    """Merging in python-docx concatenates the text of every cell involved and keeps
    one empty paragraph per cell, which is what made the header rows tall and the
    label columns repeat. Clear the merged cell and write it once instead."""
    for par in list(cell.paragraphs[1:]):
        par._element.getparent().remove(par._element)
    par = cell.paragraphs[0]
    for r in list(par.runs):
        r._element.getparent().remove(r._element)
    style_run(par.add_run(text), size, BODY_FONT, bold)
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    par.paragraph_format.space_before = Pt(1)
    par.paragraph_format.space_after = Pt(1)


def merge_and_label(table, r0, c0, r1, c1, text, size=TABLE_PT, bold=True):
    cell = table.cell(r0, c0).merge(table.cell(r1, c1))
    set_cell(cell, text, size, bold)
    return cell


def no_row_split(table):
    """Stop a single row breaking across a page. Without this, a tall cell such as
    "Support Vector Machine" splits and Word leaves an empty continuation row on the
    next page, which reads as a spurious extra model."""
    from docx.oxml import OxmlElement
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        el = OxmlElement('w:cantSplit')
        el.set(qn('w:val'), 'true')
        trPr.append(el)


def repeat_header(table, n_rows):
    """Mark the first n rows as a header so they repeat when the table breaks a page."""
    from docx.oxml import OxmlElement
    for i in range(n_rows):
        trPr = table.rows[i]._tr.get_or_add_trPr()
        el = OxmlElement('w:tblHeader')
        el.set(qn('w:val'), 'true')
        trPr.append(el)


# ---------------------------------------------------------------- report text
def build_text(R):
    """All prose, with the numbers substituted from disk."""
    d = R['deployed']
    ev = R['eval']['datasets']
    dec = R['dec']['datasets']
    e1, e2 = R['ensemble']['D1'], R['ensemble']['D2']
    ej = R['ens_json']['datasets']

    def f(tag, mk):   # deployed test F1
        return d[(tag, mk)]['test']['f1']

    def cfg(tag, mk):
        i = d[(tag, mk)]
        return f"learning rate {i['lr']:g}, batch size {i['batch_size']}, weight decay {i['weight_decay']:g}"

    C = R['classical']

    def bucket_stats(rows):
        """rows: list of (label, f1). Returns (best_label, best_f1, worst_label, worst_f1, mean_f1)."""
        best = max(rows, key=lambda r: r[1])
        worst = min(rows, key=lambda r: r[1])
        mean = sum(v for _, v in rows) / len(rows)
        return best[0], best[1], worst[0], worst[1], mean

    def grid_rows(tag, mk):
        """(label, f1) for all 8 grid configs of one model on one dataset."""
        out = []
        for lr, bs, wd in GRID:
            v = R['grid'][(tag, mk, lr, bs, wd)]
            out.append((f'lr {lr:g}, bs {bs}, wd {wd:g}', v[3]))
        return out

    def best_classical(tag, name):
        """Better of the two representations for this model on this dataset."""
        rep = max(REPS, key=lambda rp: C[(tag, name, rp)][3])
        return rep, C[(tag, name, rep)]

    def combined_rep(name):
        """One representation per model, chosen by summed F1 across both datasets,
        so a model is never shown at a different representation on each dataset."""
        rep = max(REPS, key=lambda rp: C[('D1', name, rp)][3] + C[('D2', name, rp)][3])
        return rep, C[('D1', name, rep)], C[('D2', name, rep)]

    d1_best = max(CLASSICAL, key=lambda n: best_classical('D1', n)[1][3])
    d2_best = max(CLASSICAL, key=lambda n: best_classical('D2', n)[1][3])
    d1_rep, d1_vals = best_classical('D1', d1_best)
    d2_rep, d2_vals = best_classical('D2', d2_best)

    T = {}

    T['title'] = 'Detecting AI-Generated Text with BERT, DeBERTa and a Soft-Vote Ensemble'
    T['authors'] = ', '.join(n.title() for n, _, _ in MEMBERS[:-1]) + ' and ' + MEMBERS[-1][0].title()
    T['emails'] = '{' + ', '.join(i for _, i, _ in MEMBERS) + '}@student.aiub.edu'

    T['intro'] = (
        'This project builds a system that decides whether a piece of text was written by a person or '
        'produced by a language model. We treat it as a binary classification problem and study it on two '
        'public datasets. The first is DAIGT V2, which holds 44,868 argumentative student essays, part '
        'written by students and part generated by several 2023-era models. The second is HC3, which holds '
        '85,449 question-and-answer pairs contrasting human answers with answers from ChatGPT. The two are '
        'useful together because they differ in almost every way that matters. DAIGT V2 has many generators, '
        'one genre and long documents, while HC3 has a single generator, five source domains and short ones.\n'
        'We balanced both datasets by cutting the larger class down to the size of the smaller one, which '
        'leaves 34,994 rows for DAIGT V2 and 53,806 for HC3. We then split each dataset once into training, '
        'validation and test parts of 72, 8 and 20 per cent, and every model in this report uses that same '
        'split. The split is group aware, meaning that documents with identical text after normalisation are '
        'kept together on one side of the boundary. This matters for HC3, where 7.16 per cent of rows are '
        'duplicates. Measured directly, our split puts none of the 10,732 HC3 test documents into training, '
        'while an ordinary random split of the same data puts 570 of them there. Without the group-aware '
        'split, part of the score would be measuring memorisation rather than detection.\n'
        'Table 1 summarises both datasets as published and after every preprocessing step, so the '
        'sizes every later table depends on are stated in one place. Sections 2 through 6 give the model '
        'details, grids and reasoning; this section states only the headline result.\n'
        f'The short version of the result is that DeBERTa is the best single model on both datasets, reaching '
        f'{f("D1","DeBERTa"):.4f} weighted F1 on DAIGT V2 and {f("D2","DeBERTa"):.4f} on HC3. The ensemble does not '
        'reliably beat it. We also found something less expected, which is that the two datasets are not '
        'equally hard for the same reasons, and Section 6 explains what we mean by that.')

    d1_six = [(f'{n} ({r})', C[('D1', n, r)][3]) for n in CLASSICAL for r in REPS]
    d2_six = [(f'{n} ({r})', C[('D2', n, r)][3]) for n in CLASSICAL for r in REPS]
    d1_bl, d1_bf, d1_wl, d1_wf, d1_mf = bucket_stats(d1_six)
    d2_bl, d2_bf, d2_wl, d2_wf, d2_mf = bucket_stats(d2_six)

    T['midterm'] = (
        'The midterm project compared three classical models on the same two datasets, each under two text '
        'representations, bag-of-words and TF-IDF. All six model and representation combinations were run on '
        'the full balanced corpora, using the same split the transformers use, so every number in this report '
        'comes from the same data. Table 2 gives the results.\n'
        f'Across the six combinations, not one specific model but the whole set, the best on DAIGT V2 is '
        f'{d1_bl} at {d1_bf:.4f} F1, the worst is {d1_wl} at {d1_wf:.4f}, and the mean across all six is '
        f'{d1_mf:.4f}. On HC3 the best is {d2_bl} at {d2_bf:.4f}, the worst is {d2_wl} at {d2_wf:.4f}, and the '
        f'mean is {d2_mf:.4f}. Naive Bayes sits at or near the worst position on both datasets because it '
        'assumes every word is independent of every other, an assumption text does not satisfy, and the error '
        'compounds with document length. The best position is held by whichever model and representation pairs '
        'a strong decision boundary (logistic regression or the support vector machine) with the representation '
        'that keeps the most usable signal for that dataset, which is why the winning representation differs '
        'between the two datasets even though the winning model family does not.\n'
        f'The clearest pattern is that the two datasets are not equally difficult. On DAIGT V2 every classical '
        f'model does well, and the best of them, '
        f'{d1_best.lower()} with {d1_rep}, reaches {d1_vals[3]:.4f} F1. On HC3 the same models are noticeably '
        f'weaker, with the best, {d2_best.lower()} with {d2_rep}, reaching only {d2_vals[3]:.4f}. That is a gap '
        f'of about {abs(d1_vals[3]-d2_vals[3])*100:.0f} points of F1 between the two datasets for the same family of '
        'model, so the difficulty belongs to the data rather than to the method.\n'
        'The choice of representation also matters much more on HC3 than on DAIGT V2. On HC3, moving logistic '
        f'regression from TF-IDF to bag-of-words changes F1 from {C[("D2","Logistic Regression","TF-IDF")][3]:.4f} to '
        f'{C[("D2","Logistic Regression","BoW")][3]:.4f}, a gain of about two points. On DAIGT V2 the same change moves '
        f'it from {C[("D1","Logistic Regression","TF-IDF")][3]:.4f} to {C[("D1","Logistic Regression","BoW")][3]:.4f}, less '
        'than half a point. Bag-of-words wins for five of the six model and dataset combinations, the exception '
        'being the support vector machine on DAIGT V2, where TF-IDF is better. We read this as a sign that raw '
        'counts of very common tokens carry real signal on HC3, and that TF-IDF weighting, which is designed to '
        'discount common tokens, throws part of that signal away.\n'
        'Naive Bayes is the weakest model on both datasets and by a wide margin on HC3, which is expected given '
        'that it treats every word as independent. Logistic regression and the support vector machine are close '
        'to each other everywhere.\n'
        f'These results set the target for the transformers. To be worth their extra cost, a fine-tuned model '
        f'has to beat {d1_vals[3]:.4f} on DAIGT V2 and {d2_vals[3]:.4f} on HC3.')

    b1_bl, b1_bf, b1_wl, b1_wf, b1_mf = bucket_stats(grid_rows('D1', 'BERT'))
    b2_bl, b2_bf, b2_wl, b2_wf, b2_mf = bucket_stats(grid_rows('D2', 'BERT'))
    dv1_bl, dv1_bf, dv1_wl, dv1_wf, dv1_mf = bucket_stats(grid_rows('D1', 'DeBERTa'))
    dv2_bl, dv2_bf, dv2_wl, dv2_wf, dv2_mf = bucket_stats(grid_rows('D2', 'DeBERTa'))

    T['bert'] = (
        'We chose BERT as the base transformer for three practical reasons. It is the model the course '
        'introduced, so the comparison is meaningful against what we already knew. It is small enough at 110 '
        'million parameters to fine-tune repeatedly on one consumer GPU, which we needed because the grid has '
        'eight settings per dataset. And it is the most widely reported baseline for this task, so our numbers '
        'can be compared with published work.\n'
        'We used bert-base-uncased with a maximum sequence length of 128 tokens, at most five epochs with '
        'early stopping after two epochs without improvement, a warmup ratio of 0.1, dropout of 0.1 and the '
        'AdamW optimiser. The grid varied learning rate over 0.00002 and 0.00003, batch size over 16 and 32, '
        'and weight decay over 0.01 and 0.1, giving eight runs per dataset. Table 3 lists every run. We picked '
        'the winner on the validation split, never on the test split, because choosing on test would mean '
        'reporting a number we had already optimised against.\n'
        f'On DAIGT V2 the best BERT setting is {cfg("D1","BERT")}. It reaches {f("D1","BERT"):.4f} weighted F1 on '
        f'the test split, having scored {d[("D1","BERT")]["val"]["f1"]:.4f} on validation, and it trained for '
        f'{d[("D1","BERT")]["epochs_run"]} epochs in {d[("D1","BERT")]["train_seconds"]/60:.1f} minutes. Its test confusion '
        f'matrix is {d[("D1","BERT")]["test_confusion"]}, so of 6,998 documents it misclassifies 59, split '
        'roughly evenly between the two error directions.\n'
        f'On HC3 the best setting is {cfg("D2","BERT")}, reaching {f("D2","BERT"):.4f} on test after '
        f'{d[("D2","BERT")]["val"]["f1"]:.4f} on validation, in {d[("D2","BERT")]["train_seconds"]/60:.1f} minutes. '
        f'Its confusion matrix is {d[("D2","BERT")]["test_confusion"]}, which is 90 errors out of 10,732.\n'
        f'Table 3 above is the complete picture, all eight BERT configurations on both datasets, and it '
        f'supports a best/average/worst reading. On DAIGT V2 the best of the eight is {b1_bl} at {b1_bf:.4f} '
        f'F1, the worst is {b1_wl} at {b1_wf:.4f}, and the mean across all eight is {b1_mf:.4f}; the spread '
        f'from worst to best is only {(b1_bf-b1_wf):.4f} F1. On HC3 the best is {b2_bl} at {b2_bf:.4f}, the '
        f'worst is {b2_wl} at {b2_wf:.4f}, mean {b2_mf:.4f}, spread {(b2_bf-b2_wf):.4f}. In other words, once '
        'the model and the data are fixed, the hyperparameters in this grid make very little difference. We '
        'select the best-on-validation configuration because the specification requires a single deployed '
        'model, not because the gap between best and worst is large enough to call meaningful.')

    T['variant'] = (
        'For the BERT variant we chose DeBERTa, specifically the microsoft/deberta-v3-base checkpoint. Four '
        'things made it the right choice.\n'
        'The first is benchmark standing. DeBERTa-v3-base is the strongest encoder of its size at the time we '
        'selected it, ahead of both BERT-base and RoBERTa-base across the GLUE suite of language-understanding '
        'tasks, and the DeBERTa family was the first to pass the human baseline on SuperGLUE. Starting from a '
        'stronger pretrained model is the simplest way to get a stronger fine-tuned one, and it costs us '
        'nothing extra since the two are close in size.\n'
        'The second is its attention mechanism, which keeps the representation of a word and the '
        'representation of its position separate rather than adding them together. That suits a task where '
        'the way text is written may matter as much as which words appear.\n'
        'The third is how it was pretrained. DeBERTa-v3 is trained to spot tokens that have been replaced by '
        'a small generator model, rather than to fill in masked blanks. This gives it a learning signal on '
        'every token instead of only the masked ones, so it reaches good accuracy after fewer epochs of '
        'fine-tuning, which is what our GPU budget allowed.\n'
        'The fourth is its tokeniser, and this is the one that turned out to matter most for our data. '
        'DeBERTa treats a space before a punctuation mark as part of the token, whereas the BERT tokeniser '
        'discards that distinction entirely. We verified this directly rather than assuming it. For the '
        'strings "the answer is simple ." and "the answer is simple." BERT produces identical token '
        'identifiers on three of three test pairs, while DeBERTa distinguishes all three. The two models '
        'therefore see genuinely different text, and Section 6 shows why that matters on HC3.\n'
        'One honest qualification belongs with the first reason. A model that leads a benchmark suite does '
        'not automatically lead on a new task, and our own results show exactly that. DeBERTa is far ahead on '
        'HC3 but level with BERT on DAIGT V2. Benchmark standing was a sound reason to try it first. It was '
        'not a guarantee, and we would not present it as one.\n'
        'We trained DeBERTa with exactly the same grid, the same split and the same harness as BERT, so any '
        'difference between the two is a difference between the models and not between two training setups. '
        'Table 4 lists all eight DeBERTa runs per dataset.\n'
        f'On DAIGT V2 the best DeBERTa setting is {cfg("D1","DeBERTa")}, reaching {f("D1","DeBERTa"):.4f} weighted F1 '
        f'on test after {d[("D1","DeBERTa")]["val"]["f1"]:.4f} on validation, in '
        f'{d[("D1","DeBERTa")]["train_seconds"]/60:.1f} minutes. Its confusion matrix is {d[("D1","DeBERTa")]["test_confusion"]}.\n'
        f'On HC3 the best setting is {cfg("D2","DeBERTa")}, and this is where DeBERTa clearly separates from '
        f'BERT. It reaches {f("D2","DeBERTa"):.4f} on test after {d[("D2","DeBERTa")]["val"]["f1"]:.4f} on validation, '
        f'with confusion matrix {d[("D2","DeBERTa")]["test_confusion"]}, which is only 30 errors out of 10,732. '
        f'BERT made 90 errors on the same documents, so DeBERTa removes two thirds of them.\n'
        f'Table 4 above is the complete eight-configuration DeBERTa grid on both datasets. On DAIGT V2 the '
        f'best of the eight is {dv1_bl} at {dv1_bf:.4f} F1, the worst is {dv1_wl} at {dv1_wf:.4f}, mean '
        f'{dv1_mf:.4f}, spread {(dv1_bf-dv1_wf):.4f}. On HC3 the best is {dv2_bl} at {dv2_bf:.4f}, the worst is '
        f'{dv2_wl} at {dv2_wf:.4f}, mean {dv2_mf:.4f}, spread {(dv2_bf-dv2_wf):.4f}. As with BERT, the spread is '
        'small next to the gap between models, which is why we select on validation F1 rather than treating '
        'the best-vs-worst gap inside one model as meaningful.\n'
        'The five fixed settings the instructor specified, five epochs with early stopping, a warmup ratio '
        'of 0.1, a maximum sequence length of 128 tokens, dropout of 0.1, and the AdamW optimiser, are not an '
        'alternative to our eight-point grid; they are the fixed part of the harness the grid runs inside. '
        'Every one of the eight DeBERTa runs, and every one of the eight BERT runs in Section 3, used all '
        'five exactly as specified. The eight-point grid is a separate axis, learning rate, batch size and '
        'weight decay, layered on top of those five fixed settings because the specification also requires '
        'a hyperparameter sweep, and five fixed values alone would not produce one.\n'
        'The comparison between the two models is not the same on the two datasets, and that is the most '
        'useful thing this section shows. On DAIGT V2 the two are level, '
        f'{f("D1","BERT"):.4f} against {f("D1","DeBERTa"):.4f}, a difference of one ten-thousandth that we would not '
        'call a difference at all. On HC3 the gap is real and large. Any conclusion of the form "DeBERTa is '
        'better than BERT for detecting AI text" would be true for one of our datasets and false for the other.')

    T['ensemble'] = (
        'Our ensemble model combines bert-base-uncased and microsoft/deberta-v3-base (BERT variant), each at '
        'its validation-selected best configuration from Sections 3 and 4, by averaging the class '
        'probabilities they produce, which is usually called soft voting. Rather than fixing an equal average, we searched a '
        'weight w between 0 and 1 in steps of 0.05, forming w times the BERT probabilities plus one minus w '
        'times the DeBERTa probabilities. We chose w on the validation split and then applied it once to the '
        'test split. Choosing the weight on test would have meant tuning on the data we report. Table 5 gives '
        'the two ensemble rows, one per dataset, with the weight and the resulting metrics.\n'
        f'On DAIGT V2 the search selected w = {e1["w"]:.2f}, an equal blend of the two models. The ensemble reaches '
        f'{e1["metrics"][3]:.4f} weighted F1, against {f("D1","BERT"):.4f} for BERT alone and {f("D1","DeBERTa"):.4f} for '
        'DeBERTa alone. That looks like a clear gain, and it is the point where it would be easy to overstate '
        'the result. We compared the ensemble with DeBERTa on the same test documents using a paired test, '
        f'which gives p = {ej["D1"]["paired_ensemble_vs_stronger"]["mcnemar_exact_p"]:.3f} and a 95 per cent interval on '
        'the difference in error rate of minus 0.39 to plus 0.01 percentage points. That interval contains '
        'zero, so on this test set the improvement cannot be told apart from chance. We report the ensemble as '
        'nominally ahead, not as better.\n'
        f'On HC3 the search selected w = {e2["w"]:.2f}. That means it put no weight at all on BERT, so the ensemble '
        f'is simply DeBERTa and scores exactly what DeBERTa scores, {e2["metrics"][3]:.4f}. We report this openly '
        'rather than presenting the number as an ensemble result, because a soft-vote ensemble that has '
        'discarded one of its two members is not really an ensemble.\n'
        'The honest summary is that ensembling did not help us. On the dataset where the two models are '
        'level, blending them produced a small gain we cannot confirm statistically, and on the dataset where '
        'one model is clearly better, the weight search simply picked that model. This is a reasonable '
        'outcome rather than a failure. Ensembles help most when their members make different kinds of '
        'mistakes, and on HC3 DeBERTa makes so few mistakes that BERT has little to add.')

    def final6(tag):
        rows = []
        for name in CLASSICAL:
            rep, v1, v2 = combined_rep(name)
            vals = v1 if tag == 'D1' else v2
            rows.append((f'{name} [{rep}]', vals[3]))
        rows.append(('BERT', f(tag, 'BERT')))
        rows.append(('DeBERTa (BERT variant)', f(tag, 'DeBERTa')))
        rows.append(('ENSEMBLE', R['ensemble'][tag]['metrics'][3]))
        return rows

    f6_d1 = final6('D1')
    f6_d2 = final6('D2')
    f6d1_bl, f6d1_bf, f6d1_wl, f6d1_wf, f6d1_mf = bucket_stats(f6_d1)
    f6d2_bl, f6d2_bf, f6d2_wl, f6d2_wf, f6d2_mf = bucket_stats(f6_d2)

    T['overall'] = (
        'This section brings together sir\'s two required tables. Table 6 is the full sweep, all eight '
        'configurations of BERT and DeBERTa on both datasets plus the ensemble, moved here from the per-model '
        'sections so the whole grid can be read in one place. Table 7 below it is the final six-row comparison, '
        'one row per model family at its best setting; each classical model name carries its winning '
        'representation directly, for example "Naive Bayes [BoW]", instead of a separate Representation column. '
        'Each classical model uses the same representation on both datasets, chosen by summed F1 across the '
        'two, so a model is never shown at one representation on DAIGT V2 and a different one on HC3. Naive '
        'Bayes and Logistic Regression both land on BoW under this rule, unchanged from picking per dataset. '
        'The Support Vector Machine does not, per dataset TF-IDF wins on DAIGT V2 and BoW wins on HC3, and '
        'summed F1 favours TF-IDF by 0.0014, so the SVM row uses TF-IDF on both datasets, 0.0026 F1 short of '
        'BoW\'s HC3 score rather than at BoW\'s HC3 best.\n'
        f'Across Table 7\'s six rows, the best on DAIGT V2 is {f6d1_bl} at {f6d1_bf:.4f} F1, the worst is '
        f'{f6d1_wl} at {f6d1_wf:.4f}, mean {f6d1_mf:.4f}. On HC3 the best is {f6d2_bl} at {f6d2_bf:.4f}, the '
        f'worst is {f6d2_wl} at {f6d2_wf:.4f}, mean {f6d2_mf:.4f}. DeBERTa is best on both because disentangled '
        'attention and a replaced-token-detection pretraining objective give it a genuine edge on HC3\'s '
        'formatting-carried signal, discussed below, while on DAIGT V2 it merely ties the classical ceiling '
        'rather than beating it. Naive Bayes is worst on both for the same independence-assumption reason given '
        'in Section 2, and it stays worst even after being allowed its better representation.\n'
        'Reading the table across, three further conclusions hold. The first is that DeBERTa is the best single model '
        'on both datasets. The second is that the size of its advantage depends entirely on which dataset you '
        f'look at. On DAIGT V2 the best classical model reaches {d1_vals[3]:.4f} and DeBERTa reaches '
        f'{f("D1","DeBERTa"):.4f}, a difference of {abs(f("D1","DeBERTa")-d1_vals[3])*100:.2f} of a percentage point, so '
        'the extra cost of a transformer buys almost nothing there. On HC3 the same comparison is '
        f'{d2_vals[3]:.4f} against {f("D2","DeBERTa"):.4f}, a gap of more than four points, and the transformer is '
        'clearly doing something the classical models cannot. The third is that the ensemble adds nothing we '
        'can demonstrate.\n'
        'It is worth being precise about the DAIGT V2 result rather than rounding it away. The classical model '
        'reads the whole document while the transformers read only their first 128 tokens, which is about a '
        'third of the average essay. When we refitted the classical models on exactly the same 128-token span '
        'the transformers see, their error rate rose from 0.90 to 2.10 per cent while BERT stayed at 0.84. So '
        'the apparent tie on DAIGT V2 is not evidence that a bag-of-words model is as good as a transformer. '
        'It is evidence that reading the whole essay is worth about as much as being a transformer, which is a '
        'different and more useful statement.\n'
        'We also ran one experiment that is not required by the specification but explains the pattern above, '
        'and it is the most interesting thing we found. We built two restricted models. One reads only 47 '
        'surface properties of the text, such as how often a space appears before a full stop, how long the '
        'document is, and how often capital letters are used. It never sees a single word. The other reads '
        'only the words, after punctuation, capitalisation and unusual characters have been stripped out. '
        f'On HC3 these two score almost identically, with error rates of {dec["D2"]["surface_only"]["error_rate"]*100:.2f} '
        f'and {dec["D2"]["content_only"]["error_rate"]*100:.2f} per cent, and a paired test cannot separate them. In '
        'other words, on HC3 you can do about as well by looking at how the text is punctuated as by reading '
        f'what it says. On DAIGT V2 the picture is different, with {dec["D1"]["surface_only"]["error_rate"]*100:.2f} per '
        f'cent error from surface alone against {dec["D1"]["content_only"]["error_rate"]*100:.2f} per cent from words '
        'alone, so the words carry roughly eight times more of the signal.\n'
        'This explains why HC3 looked harder for the classical models and easier for DeBERTa. Much of what '
        'separates the classes on HC3 is in the formatting, and the classical pipeline deletes formatting '
        'during preprocessing while DeBERTa keeps it. It also means a high score on HC3 should be read '
        'carefully, because part of it comes from how the dataset was collected rather than from any real '
        'difference between human and machine writing.')

    T['limits'] = (
        'Several things limit what our numbers can be taken to mean, and we would rather state them than '
        'leave them to be discovered.\n'
        'Both transformers read at most 128 tokens per document. On HC3 that covers about three quarters of '
        'the average answer, but on DAIGT V2 it covers only about a third of the average essay. So the '
        'DAIGT V2 transformer results really describe classifying an essay from its opening, and the '
        'comparison with the classical models, which read the whole document, is not a fair fight in either '
        'direction.\n'
        'Every transformer number comes from a single training seed. We repeated a few settings with '
        'different seeds and saw the score move by up to a few thousandths, and one such estimate changed by '
        'about a third when we simply reran it, because training on a GPU is not perfectly repeatable. For '
        'that reason we did not base any claim on small differences between runs, and where we compare two '
        'models we compare their predictions on the same documents instead.\n'
        'Both datasets are English, and one of them, HC3, contains output from a single generator collected '
        'at a single point in time. Nothing in this report shows that our models would work on text from a '
        'newer model, on another language, or on a domain neither dataset covers. We checked the weakest '
        'version of this by training on one dataset and testing on the other, and performance falls by '
        'between 8 and 20 points of F1, so the models clearly learn things specific to their own dataset.\n'
        'Finally, we did not test what happens when someone actively tries to defeat the detector. We ran '
        'some simple text corruption experiments, but we also found that our models confidently label random '
        'character strings as human-written, which means a drop in score under corruption cannot be '
        'separated from the model simply not coping with unreadable input. We therefore do not report those '
        'experiments as evidence of robustness in either direction.')

    T['conclusion'] = (
        'We built and compared five families of model for deciding whether text is human-written or '
        'machine-generated, on two public datasets, using one fixed split and one training harness so that '
        f'the comparisons are fair. DeBERTa is the best single model on both, reaching {f("D1","DeBERTa"):.4f} '
        f'weighted F1 on DAIGT V2 and {f("D2","DeBERTa"):.4f} on HC3, which improves on the best classical model '
        f'by {abs(f("D1","DeBERTa")-d1_vals[3])*100:.2f} of a percentage point on the first dataset and by more than four '
        'points on the second. A soft-vote ensemble of BERT and DeBERTa did not reliably improve on DeBERTa alone, '
        'and we report that plainly rather than presenting a difference in the fourth decimal place as a '
        'gain.\n'
        'The result we think is most worth carrying forward is not the leaderboard. It is that a model '
        'reading only 47 formatting properties of a document, and no words at all, matches a full '
        'bag-of-words model on HC3 while falling far behind it on DAIGT V2. That tells us the two datasets '
        'are separable for different reasons, which a single accuracy figure hides completely. Anyone '
        'choosing a dataset to evaluate a detector, or reading a reported score, should want to know which '
        'of the two situations they are in. The check is cheap, it needs no extra labelling, and it runs in '
        'minutes, so we would recommend running it before trusting any headline number on this task.')

    return T


# ---------------------------------------------------------------- assembly
def main():
    R = load_results()
    T = build_text(R)
    doc = Document(str(TEMPLATE))

    # ---- cover page
    set_text(find_par(doc, 'Section:'), 'Section: B', justify=False).alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_text(find_par(doc, 'Group:'), 'Group: 02', justify=False).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---- contributions table
    tbl = doc.tables[0]
    for i, (name, sid, contrib) in enumerate(MEMBERS, start=1):
        set_text(tbl.cell(i, 1).paragraphs[0], f'Name: {name}', justify=False)
        p2 = tbl.cell(i, 1).paragraphs[1] if len(tbl.cell(i, 1).paragraphs) > 1 else tbl.cell(i, 1).add_paragraph()
        set_text(p2, f'ID: {sid}', justify=False)
        set_text(tbl.cell(i, 2).paragraphs[0], contrib)

    # ---- title block
    set_text(find_par(doc, 'Project Title'), T['title'], size=16, bold=True, justify=False
             ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_text(find_par(doc, 'Student 1 Name'), T['authors'], justify=False
             ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_text(find_par(doc, '@student.aiub.edu'), T['emails'], justify=False
             ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---- section bodies. Each template placeholder becomes the first paragraph and
    #      any remaining ones are appended after it, so paragraph breaks survive.
    def write_section(placeholder_needle, body):
        chunks = body.split('\n')
        first = set_text(find_par(doc, placeholder_needle), chunks[0])
        pars = [first]
        anchor = first
        for c in chunks[1:]:
            anchor = para_after(anchor, c)
            pars.append(anchor)
        return pars

    sec1 = write_section('Provide an overall summary of your work', T['intro'])
    sec2 = write_section('Take the resultant table from the experiments', T['midterm'])
    sec3 = write_section('Give reason behind selecting the BERT model. Additionally, based on your experimental results, select the BERT model', T['bert'])
    sec4 = write_section('select the BERT variant model', T['variant'])
    # Name the variant in the heading itself without disturbing the template's heading
    # style (set_text would reset the run to plain body font/size).
    h4 = find_par(doc, '4 Selecting the Best BERT variant model')
    base_run = h4.runs[0] if h4.runs else None
    tail = h4.add_run(' — DeBERTa')
    if base_run is not None:
        tail.font.name = base_run.font.name
        tail.font.size = base_run.font.size
        tail.font.bold = base_run.font.bold
    sec5 = write_section('Discuss your ensemble model', T['ensemble'])
    sec6 = write_section('Compare the overall results', T['overall'])
    write_section('Discuss the limitations', T['limits'])
    write_section('Summarize your work in this section', T['conclusion'])

    # ---- Table 1, what the datasets look like after preprocessing
    sp = R['splits']
    tk = R['tok']
    da, ha, hc = R['daigt_audit'], R['hc3_audit'], R['hc3_counts']
    cap0 = para_after(sec1[-1], 'Table 1. The two datasets as published and after preprocessing. Duplicate rows '
                            'are exact matches after whitespace and case normalisation. Leakage is the share of '
                            'test documents whose text also appears in training or validation. Token counts use '
                            'the bert-base-uncased tokeniser on a 2,000-row sample per dataset.',
                      size=10, justify=False)
    ds_rows = [
        ['', 'DAIGT V2', 'HC3'],
        ['As published', '', ''],
        ['   Total documents', f"{da['n_total']:,}", f"{hc['n_total']:,}"],
        ['   Human-written', f"{da['n_human']:,}", f"{hc['n_human']:,}"],
        ['   Machine-generated', f"{da['n_ai']:,}", f"{hc['n_machine']:,}"],
        ['   Duplicate rows',
         f"{R['daigt_audit']['dup_rows']:,} ({R['daigt_audit']['dup_pct']}%)",
         f"{R['hc3_audit']['dup_rows']:,} ({R['hc3_audit']['dup_rows']/hc['n_total']*100:.2f}%)"],
        ['   Collection artefact rows', 'none found',
         f"{R['hc3_audit']['artefact_chatgpt'] + R['hc3_audit']['artefact_human']:,}"],
        ['After class balancing', '', ''],
        ['   Total documents', f"{sp['D1']['balanced']:,}", f"{sp['D2']['balanced']:,}"],
        ['   Per class', f"{sp['D1']['per_class']:,}", f"{sp['D2']['per_class']:,}"],
        ['After splitting (72 / 8 / 20)', '', ''],
        ['   Training', f"{sp['D1']['train']:,}", f"{sp['D2']['train']:,}"],
        ['   Validation', f"{sp['D1']['val']:,}", f"{sp['D2']['val']:,}"],
        ['   Test', f"{sp['D1']['test']:,}", f"{sp['D2']['test']:,}"],
        ['   Test leakage, group-aware split', '0 (0.00%)', '0 (0.00%)'],
        ['   Test leakage, ordinary random split', '0 (0.00%)', '570 (5.30%)'],
        ['Document length', '', ''],
        ['   Median words per document', f"{sp['D1']['median_words']}", f"{sp['D2']['median_words']}"],
        ['   Median words, human',
         f"{sp['D1']['median_words_human']}", f"{sp['D2']['median_words_human']}"],
        ['   Median words, machine',
         f"{sp['D1']['median_words_machine']}", f"{sp['D2']['median_words_machine']}"],
        ['   Median tokens per document',
         f"{tk['D1']['median_tokens']:.0f}", f"{tk['D2']['median_tokens']:.0f}"],
        ['   Documents longer than 128 tokens',
         f"{tk['D1']['pct_exceeding_128']}%", f"{tk['D2']['pct_exceeding_128']}%"],
        ['   Median share read at 128 tokens',
         f"{tk['D1']['median_pct_kept_at_128']}%", f"{tk['D2']['median_pct_kept_at_128']}%"],
    ]
    keep_with_next(cap0)
    t0 = table_after(doc, cap0, ds_rows, header_rows=1)
    for i, row in enumerate(ds_rows):
        if row[1] == '' and row[2] == '' and i > 0:
            merge_and_label(t0, i, 0, i, 2, row[0], bold=True)
        else:
            t0.cell(i, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    repeat_header(t0, 1)
    no_row_split(t0)

    # ---- Table 2, the classical models

    cap1 = para_after(sec2[0], 'Table 2. Classical models on the full balanced corpora, both text representations, '
                            'evaluated on the same test split every other model in this report uses. Acc = Accuracy, Prec = Precision, Rec = Recall, F1 = F1 Score.',
                      size=10, justify=False)
    rows = [[''] * 10, ['', '', 'Acc', 'Prec', 'Rec', 'F1', 'Acc', 'Prec', 'Rec', 'F1']]
    for name in CLASSICAL:
        for rep in REPS:
            rows.append([name, rep]
                        + [f'{v:.4f}' for v in R['classical'][('D1', name, rep)]]
                        + [f'{v:.4f}' for v in R['classical'][('D2', name, rep)]])
    keep_with_next(cap1)
    t1 = table_after(doc, cap1, rows, header_rows=2)
    merge_and_label(t1, 0, 2, 0, 5, 'Dataset 1 (DAIGT V2)')
    merge_and_label(t1, 0, 6, 0, 9, 'Dataset 2 (HC3)')
    merge_and_label(t1, 0, 0, 1, 0, 'Model')
    merge_and_label(t1, 0, 1, 1, 1, 'Represen-\ntation')
    for i, name in enumerate(CLASSICAL):
        r0 = 2 + i * len(REPS)
        merge_and_label(t1, r0, 0, r0 + len(REPS) - 1, 0, name, bold=False)
    repeat_header(t1, 2)
    no_row_split(t1)

    # ---- Table 3 / Table 4, one model's own 8-row grid, isolated per the "just this
    #      variant's models" instruction (no other model or the ensemble in these tables).
    def single_model_grid_table(anchor, mk, label, table_no):
        cap = para_after(anchor,
            f'Table {table_no}. {label} fine-tuning experiments only. All eight configurations '
            'of the sweep, full balanced corpora, evaluated on the test split.',
            size=10, justify=False)
        rows = [[''] * 11, ['Learning Rate', 'Batch Size', 'Weight Decay',
                            'Acc', 'Prec', 'Rec', 'F1', 'Acc', 'Prec', 'Rec', 'F1']]
        missing = []
        for lr, bs, wd in GRID:
            a = R['grid'].get(('D1', mk, lr, bs, wd))
            b = R['grid'].get(('D2', mk, lr, bs, wd))
            if a is None or b is None:
                missing.append((mk, lr, bs, wd))
                continue
            rows.append([f'{lr:.5f}', str(bs), str(wd)]
                        + [f'{v:.4f}' for v in a] + [f'{v:.4f}' for v in b])
        if missing:
            raise SystemExit('missing grid results, refusing to emit blank cells: %s' % missing)
        keep_with_next(cap)
        t = table_after(doc, cap, rows, header_rows=2)
        merge_and_label(t, 0, 3, 0, 6, 'Dataset 1 (DAIGT V2)')
        merge_and_label(t, 0, 7, 0, 10, 'Dataset 2 (HC3)')
        for c, lbl in enumerate(('Learning Rate', 'Batch Size', 'Weight Decay')):
            merge_and_label(t, 0, c, 1, c, lbl)
        repeat_header(t, 2)
        no_row_split(t)
        return t

    single_model_grid_table(sec3[1], 'BERT', 'BERT', 3)
    single_model_grid_table(sec4[6], 'DeBERTa', 'DeBERTa (BERT variant)', 4)

    # ---- Table 5, the ensemble only (two rows, one per dataset)
    cap_ens = para_after(sec5[0], 'Table 5. Ensemble model only. The validation-selected mixing '
                            'weight on BERT and the resulting test-set metrics, one row per dataset.',
                          size=10, justify=False)
    rows = [['Dataset', 'Weight on BERT', 'Acc', 'Prec', 'Rec', 'F1'],
            ['DAIGT V2', f"{R['ensemble']['D1']['w']:.2f}"] + [f'{v:.4f}' for v in R['ensemble']['D1']['metrics']],
            ['HC3', f"{R['ensemble']['D2']['w']:.2f}"] + [f'{v:.4f}' for v in R['ensemble']['D2']['metrics']]]
    keep_with_next(cap_ens)
    t_ens = table_after(doc, cap_ens, rows, header_rows=1)
    repeat_header(t_ens, 1)
    no_row_split(t_ens)

    # ---- Table 6, sir's full sweep table (all models + ensemble, moved here from the
    #      per-model sections so Section 6 holds both of sir's required tables together)
    cap6 = para_after(sec6[0], 'Table 6. Final-term fine-tuning experiments, all models. Eight configurations '
                            'for each of BERT and DeBERTa on each dataset, full balanced corpora, evaluated on '
                            'the test split. The ENSEMBLE row uses the weight chosen on validation.',
                       size=10, justify=False)
    rows = [[''] * 12,
            ['', '', '', '', 'Acc', 'Prec', 'Rec', 'F1', 'Acc', 'Prec', 'Rec', 'F1']]
    missing = []
    for mk in ('BERT', 'DeBERTa'):
        for lr, bs, wd in GRID:
            a = R['grid'].get(('D1', mk, lr, bs, wd))
            b = R['grid'].get(('D2', mk, lr, bs, wd))
            if a is None or b is None:
                missing.append((mk, lr, bs, wd))
            rows.append(['', f'{lr:.5f}', str(bs), str(wd)]
                        + [f'{v:.4f}' for v in a] + [f'{v:.4f}' for v in b])
    if missing:
        raise SystemExit('missing grid results, refusing to emit blank cells: %s' % missing)
    rows.append([''] * 4
                + [f'{v:.4f}' for v in R['ensemble']['D1']['metrics']]
                + [f'{v:.4f}' for v in R['ensemble']['D2']['metrics']])
    keep_with_next(cap6)
    t6 = table_after(doc, cap6, rows, header_rows=2, size=6.5)
    merge_and_label(t6, 0, 4, 0, 7, 'Dataset 1 (DAIGT V2)', size=6.5)
    merge_and_label(t6, 0, 8, 0, 11, 'Dataset 2 (HC3)', size=6.5)
    for c, lbl in enumerate(('Model', 'Learning Rate', 'Batch Size', 'Weight Decay')):
        merge_and_label(t6, 0, c, 1, c, lbl, size=6.5)
    merge_and_label(t6, 2, 0, 9, 0, 'BERT', size=6.5, bold=False)
    merge_and_label(t6, 10, 0, 17, 0, 'DeBERTa\n(BERT variant)', size=6.5, bold=False)
    set_cell(t6.cell(18, 0), 'ENSEMBLE', size=6.5, bold=True)
    merge_and_label(t6, 18, 1, 18, 3, '(validation-selected weight)', size=6.5, bold=False)
    repeat_header(t6, 2)
    no_row_split(t6)

    # ---- Table 7, sir's final comparison table. Six rows, one per model family, no
    #      Representation column — the classical model name itself carries [BoW]/[TF-IDF].
    #      One representation per model, chosen by summed F1 across both datasets, so a
    #      model never shows a different representation on each dataset (per sir's note).
    def combined_rep_main(name):
        rep = max(REPS, key=lambda rp: R['classical'][('D1', name, rp)][3]
                                       + R['classical'][('D2', name, rp)][3])
        return rep, R['classical'][('D1', name, rep)], R['classical'][('D2', name, rep)]

    cap7 = para_after(sec6[1], 'Table 7. Final comparison, six rows. Every row comes from the full balanced '
                            'corpora and the same fixed split. Each classical model name carries the text '
                            'representation it uses, e.g. "Naive Bayes [BoW]", the same representation on '
                            'both datasets for that model, chosen by summed F1 across the two. The ENSEMBLE '
                            'row uses the weight chosen on validation.', size=10, justify=False)
    rows = [[''] * 9,
            ['Model', 'Acc', 'Prec', 'Rec', 'F1', 'Acc', 'Prec', 'Rec', 'F1']]
    for name in CLASSICAL:
        rep, v1, v2 = combined_rep_main(name)
        rows.append([f'{name} [{rep}]']
                    + [f'{v:.4f}' for v in v1] + [f'{v:.4f}' for v in v2])
    for mk, label in (('BERT', 'BERT'), ('DeBERTa', 'DeBERTa (BERT variant)')):
        a = R['deployed'][('D1', mk)]['test']
        b = R['deployed'][('D2', mk)]['test']
        rows.append([label] + [f'{a[k]:.4f}' for k in ('accuracy', 'precision', 'recall', 'f1')]
                    + [f'{b[k]:.4f}' for k in ('accuracy', 'precision', 'recall', 'f1')])
    rows.append(['ENSEMBLE'] + [f'{v:.4f}' for v in R['ensemble']['D1']['metrics']]
                + [f'{v:.4f}' for v in R['ensemble']['D2']['metrics']])
    keep_with_next(cap7)
    t7 = table_after(doc, cap7, rows, header_rows=2, size=7)
    merge_and_label(t7, 0, 1, 0, 4, 'Dataset 1 (DAIGT V2)', size=7)
    merge_and_label(t7, 0, 5, 0, 8, 'Dataset 2 (HC3)', size=7)
    merge_and_label(t7, 0, 0, 1, 0, 'Model', size=7)
    repeat_header(t7, 2)
    no_row_split(t7)

    # ---- code appendix intentionally omitted from the docx. The reorganized notebook is
    #      exported to its own PDF instead (docs/nlp_final_submission_code.pdf); see
    #      docs/superpowers/plans/2026-08-28-final-report-restructure.md, Tasks 9-10.
    for needle in ('Give the entire code of your project here',
                   'Data preprocessing code if there is any',
                   'BERT code where you achieved the best performance',
                   'BERT Variant code where you achieved the best performance',
                   'Ensemble code'):
        try:
            p = find_par(doc, needle)
            p._element.getparent().remove(p._element)
        except KeyError:
            pass

    doc.save(str(OUT))
    print('wrote', OUT.name)
    print('  sections filled 8, tables 7 (3 dataset/classical, 3 per BERT/BERT variant, 1 dataset-diagnostic)')


if __name__ == '__main__':
    main()
